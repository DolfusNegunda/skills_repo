
"""
generate_standard_doc.py

Generic, organization-agnostic document generator for standardized process
documents (starting with Lessons Learned). Works against ANY branded .docx
base template, as long as that template follows the standardized section
skeleton (see references/data-schema.md) and uses {{TOKEN}} placeholders
for single-value fields.

Design principles (what makes this portable across companies):
  1. Singular fields (project name, code, dates, org name, etc.) are plain
     {{TOKEN}} placeholders replaced via find/replace across paragraphs,
     table cells, headers/footers, AND Word content controls (w:sdt) if
     the org's real template happens to use the built-in Cover Page
     gallery. No positional/index assumptions.
  2. Repeating list sections (team roster, deliverables, positives,
     challenges, learnings) are located by the STANDARD HEADING TEXT that
     precedes them, not by table index. This means the base template can
     have extra paragraphs, different table ordering, or additional
     sections inserted, and the script still finds the right table.
  3. The org's own branding (logo, colors, fonts, footer/legal text) lives
     entirely inside the base template .docx supplied by the caller. This
     script never assumes a specific company's branding.

Usage:
    python generate_standard_doc.py --data project_data.json \
        --template org_branded_template.docx \
        --logo org_logo.png \
        --output "Lessons_Learned.docx"

Or import generate() directly from Python.
"""
import argparse
import copy
import io
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from PIL import Image


# ---------------------------------------------------------------------------
# Generic {{TOKEN}} replacement (paragraphs, table cells, headers/footers,
# and Word content controls / sdt elements)
# ---------------------------------------------------------------------------

def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """Replace `old` with `new` across a paragraph's runs, merging as needed.
    Preserves the formatting of the first run in the matched span."""
    runs = paragraph.runs
    full_text = "".join(r.text for r in runs)
    start = full_text.find(old)
    if start == -1:
        return False
    end = start + len(old)

    positions, pos = [], 0
    for r in runs:
        positions.append((pos, pos + len(r.text)))
        pos += len(r.text)

    touched_idx = [i for i, (s, e) in enumerate(positions) if e > start and s < end]
    if not touched_idx:
        return False

    first_i, last_i = touched_idx[0], touched_idx[-1]
    first_run, last_run = runs[first_i], runs[last_i]
    prefix = first_run.text[: start - positions[first_i][0]]
    suffix = last_run.text[end - positions[last_i][0]:]

    if first_i == last_i:
        first_run.text = prefix + new + suffix
    else:
        first_run.text = prefix + new
        last_run.text = suffix
        for i in range(first_i + 1, last_i):
            runs[i].text = ""
    return True


def replace_token_everywhere(doc: Document, token: str, value: str) -> int:
    """Replace every occurrence of `token` (e.g. "{{PROJECT_CODE}}") with
    `value`, across the main body, tables (including nested tables), and
    headers/footers of every section. Also handles Word content controls
    (w:sdt), which live outside doc.paragraphs (used by templates built
    from Word's "Cover Pages" gallery)."""
    count = 0

    def scan_paragraphs(paragraphs):
        nonlocal count
        for p in paragraphs:
            while replace_in_paragraph(p, token, value):
                count += 1

    def scan_tables(tables):
        nonlocal count
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    scan_paragraphs(cell.paragraphs)
                    scan_tables(cell.tables)  # nested tables, if any

    scan_paragraphs(doc.paragraphs)
    scan_tables(doc.tables)

    for section in doc.sections:
        for hf in (section.header, section.footer, section.first_page_header,
                   section.first_page_footer, section.even_page_header,
                   section.even_page_footer):
            scan_paragraphs(hf.paragraphs)
            scan_tables(hf.tables)

    # Word content controls (sdt) — e.g. templates built from the "Cover
    # Pages" building-block gallery store title/date text here, outside
    # doc.paragraphs entirely.
    for sdt in doc.element.body.iter(qn("w:sdt")):
        t_nodes = sdt.findall(".//" + qn("w:t"))
        joined = "".join(t.text or "" for t in t_nodes)
        if token in joined and t_nodes:
            new_text = joined.replace(token, value)
            t_nodes[0].text = new_text
            for t in t_nodes[1:]:
                t.text = ""
            count += 1

    return count


def replace_tokens(doc: Document, mapping: dict) -> None:
    """Replace every {{TOKEN}} -> value pair in `mapping` throughout the doc."""
    for token, value in mapping.items():
        replace_token_everywhere(doc, token, str(value))


# ---------------------------------------------------------------------------
# Locating repeating-list tables by their standard heading text (NOT by
# table index — this is what makes the script portable across different
# organizations' base templates)
# ---------------------------------------------------------------------------

def find_table_after_heading(doc: Document, heading_text: str):
    """Return the first <w:tbl> element that appears anywhere after a
    paragraph whose text matches `heading_text` (case-insensitive,
    substring match), scanning the document body in document order."""
    body = doc.element.body
    found_heading = False
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = "".join(
                node.text or "" for node in child.iter(qn("w:t"))
            )
            if heading_text.lower() in text.lower():
                found_heading = True
        elif tag == "tbl" and found_heading:
            from docx.table import Table
            return Table(child, doc)
    raise ValueError(f'No table found after a heading matching "{heading_text}"')


# ---------------------------------------------------------------------------
# Table / cell helpers
# ---------------------------------------------------------------------------

def clone_row(table, row_idx: int):
    """Deep-clone table.rows[row_idx] and insert it immediately after, to
    preserve cell borders/shading/fonts. Returns the new row object."""
    src_tr = table.rows[row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    src_tr.addnext(new_tr)
    return table.rows[row_idx + 1]


def set_cell_text(cell, text: str) -> None:
    """Replace a cell's visible text entirely with `text`, preserving the
    first paragraph/run's formatting. Cells sometimes hold multiple
    paragraphs (multi-line comments) — all but the first are removed so no
    leftover template text survives."""
    paragraphs = cell.paragraphs
    first_p = paragraphs[0]
    if first_p.runs:
        first_p.runs[0].text = text
        for r in first_p.runs[1:]:
            r.text = ""
    else:
        first_p.add_run(text)
    for p in paragraphs[1:]:
        p._p.getparent().remove(p._p)


def remove_extra_rows(table, keep_through_idx: int) -> None:
    """Delete every row after index `keep_through_idx` (0-based, inclusive)."""
    while len(table.rows) > keep_through_idx + 1:
        table.rows[keep_through_idx + 1]._tr.getparent().remove(
            table.rows[keep_through_idx + 1]._tr
        )


def fill_repeating_table(table, header_row_count: int, rows: list, column_getters: list) -> None:
    """Fill a repeating-list table (one header row or more, followed by ONE
    template data row in the base template) with `rows` (a list of dicts).
    `column_getters` is a list of callables, one per physical column, each
    mapping a row dict -> the string to put in that cell."""
    for i, item in enumerate(rows):
        row = table.rows[header_row_count] if i == 0 else clone_row(table, header_row_count + i - 1)
        for col_idx, getter in enumerate(column_getters):
            if col_idx < len(row.cells):
                set_cell_text(row.cells[col_idx], getter(item))
    remove_extra_rows(table, header_row_count + len(rows) - 1)


def replace_table_with_skeleton(doc: Document, table, skeleton_xml_path: str):
    """Swap `table` for the stored skeleton (used when a document format
    choice — e.g. deliverables format A vs B — requires a differently
    shaped table than whatever the base template shipped with)."""
    old_tbl = table._tbl
    with open(skeleton_xml_path, "rb") as f:
        new_tbl = parse_xml(f.read())
    old_tbl.addprevious(new_tbl)
    old_tbl.getparent().remove(old_tbl)
    from docx.table import Table
    return Table(new_tbl, doc)


# ---------------------------------------------------------------------------
# Logo swap (organization/project logo lives at word/media/image1.*, i.e.
# whichever image appears first in the template's header)
# ---------------------------------------------------------------------------

def swap_logo(docx_path: str, new_logo_path: str) -> None:
    """Replace the organization/project logo image (word/media/image1.*)
    in-place, preserving the required file extension."""
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        image1_candidates = [n for n in names if n.startswith("word/media/image1.")]
        if not image1_candidates:
            raise FileNotFoundError(
                "No word/media/image1.* found in the template — the base "
                "template must contain at least one image for the logo "
                "swap to have something to replace."
            )
        image1_name = image1_candidates[0]
    ext_needed = image1_name.rsplit(".", 1)[-1].lower()

    im = Image.open(new_logo_path)
    buf = io.BytesIO()
    fmt = "JPEG" if ext_needed in ("jpg", "jpeg") else "PNG"
    if fmt == "JPEG" and im.mode in ("RGBA", "P"):
        im = im.convert("RGB")
    im.save(buf, format=fmt)
    new_bytes = buf.getvalue()

    tmp_path = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == image1_name:
                data = new_bytes
            zout.writestr(item, data)
    Path(tmp_path).replace(docx_path)


# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------

def generate(
    data: dict,
    template_path: str,
    output_path: str,
    logo_path: str | None = None,
    assets_dir: str = ".",
) -> str:
    """Generate a standardized Lessons Learned .docx using ANY organization's
    branded base template.

    `data` keys — see references/data-schema.md for the authoritative list:
      org_name, project_code, project_name, project_manager, cover_date,
      metadata_date, team (list of {name, position, allocation?}),
      deliverables_format ("A" or "B"), deliverables (list matching format),
      positives (list of {title, comment}),
      challenges (list of {title, comment}),
      learnings (list of {title, actions}).

    The base template must contain the standard section headings verbatim:
    "Project Team", "Deliverables", "Positives", "Challenges", "Learnings
    and Suggestions" — and single-value placeholders as {{TOKENS}}:
    {{ORG_NAME}}, {{PROJECT_CODE}}, {{PROJECT_NAME}}, {{PROJECT_MANAGER}},
    {{COVER_DATE}}, {{METADATA_DATE}}.
    """
    import shutil
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # --- Singular fields: generic token replacement, works regardless of
    #     whether the org's template uses plain paragraphs or Word content
    #     controls for its cover page ---
    replace_tokens(doc, {
        "{{ORG_NAME}}": data.get("org_name", ""),
        "{{PROJECT_CODE}}": data["project_code"],
        "{{PROJECT_NAME}}": data["project_name"],
        "{{PROJECT_MANAGER}}": data["project_manager"],
        "{{COVER_DATE}}": data.get("cover_date", data.get("metadata_date", "")),
        "{{METADATA_DATE}}": data.get("metadata_date", data.get("cover_date", "")),
    })

    # --- Project Team: located by heading, not table index ---
    t_team = find_table_after_heading(doc, "Project Team")
    has_allocation = len(t_team.rows[0].cells) >= 3 and any(
        m.get("allocation") for m in data["team"]
    )
    getters = [lambda m: m["name"], lambda m: m["position"]]
    if has_allocation:
        getters.append(lambda m: m.get("allocation", ""))
    fill_repeating_table(t_team, header_row_count=1, rows=data["team"], column_getters=getters)

    # --- Deliverables: located by heading; format A/B skeleton swapped in
    #     first so the correct table shape is guaranteed regardless of what
    #     shape the base template's own deliverables table happens to be ---
    fmt = data.get("deliverables_format", "B")
    t_deliv = find_table_after_heading(doc, "Deliverables")
    skeleton_filename = f"deliverables-table-format-{fmt}.xml"
    skeleton_path = str(Path(assets_dir) / skeleton_filename)
    t_deliv = replace_table_with_skeleton(doc, t_deliv, skeleton_path)
    if fmt == "A":
        def desc_getter(item):
            description = item["description"]
            status_note = item.get("status_note")
            if status_note:
                description = f"{description}\n\u2022 {status_note}"
            return description
        getters = [lambda item: item["item"], desc_getter]
    else:
        getters = [
            lambda item: item["document"],
            lambda item: item.get("status", ""),
            lambda item: item.get("notes", ""),
        ]
    fill_repeating_table(t_deliv, header_row_count=1, rows=data["deliverables"], column_getters=getters)

    # --- Positives ---
    t_pos = find_table_after_heading(doc, "Positives")
    fill_repeating_table(
        t_pos, header_row_count=1, rows=data["positives"],
        column_getters=[lambda i: i["title"], lambda i: i["comment"]],
    )

    # --- Challenges ---
    t_chal = find_table_after_heading(doc, "Challenges")
    fill_repeating_table(
        t_chal, header_row_count=1, rows=data["challenges"],
        column_getters=[lambda i: i["title"], lambda i: i["comment"]],
    )

    # --- Learnings and Suggestions ---
    t_learn = find_table_after_heading(doc, "Learnings and Suggestions")
    fill_repeating_table(
        t_learn, header_row_count=1, rows=data["learnings"],
        column_getters=[lambda i: i["title"], lambda i: i["actions"]],
    )

    doc.save(output_path)

    if logo_path:
        swap_logo(output_path, logo_path)

    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate a standardized Lessons Learned document.")
    parser.add_argument("--data", required=True, help="Path to a JSON file matching the data schema.")
    parser.add_argument("--template", required=True, help="Path to the organization's branded base template .docx")
    parser.add_argument("--logo", default=None, help="Path to the organization/project logo image (optional).")
    parser.add_argument("--output", required=True, help="Output .docx path.")
    parser.add_argument("--assets-dir", default=".", help="Directory containing the deliverables table skeleton XML files.")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    out = generate(data, args.template, args.output, args.logo, args.assets_dir)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
