"""Generate the example client file (client_qbr.docx) used by examples/README.md.

Kept as a generator (not a committed binary) so the example is reproducible and
never goes stale. Run from this examples/ directory:  python make_example_docx.py
"""
from pathlib import Path
from docx import Document

doc = Document()
doc.add_heading("Quarterly Business Review", level=0)

for label, value in [("Client", "Globex Corporation"), ("Prepared by", "Jane Smith"),
                     ("Period", "Q3 2026"), ("Date", "2026-07-13")]:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)

doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "This document summarizes the quarterly performance for Globex Corporation. "
    "It is confidential and prepared solely for the named client. All figures are "
    "unaudited and subject to revision under our standard engagement terms."
)
doc.add_paragraph("Globex Corporation continued strong momentum this period.")

doc.add_heading("Key Achievements", level=1)
for item in ["Launched the new analytics portal", "Reduced churn by 12%",
             "Signed 3 enterprise accounts"]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Financials", level=1)
table = doc.add_table(rows=2, cols=2)
table.style = "Light Grid Accent 1"
table.cell(0, 0).text = "Revenue"
table.cell(0, 1).text = "$4.2M"
table.cell(1, 0).text = "Net New ARR"
table.cell(1, 1).text = "$820K"

out = Path(__file__).parent / "client_qbr.docx"
doc.save(out)
print(f"wrote {out}")
