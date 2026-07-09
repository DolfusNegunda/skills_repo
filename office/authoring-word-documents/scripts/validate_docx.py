"""Validate an authored .docx before it ships.

Deterministic quality gate for a produced Word document. Emits a machine-readable
JSON report and exits non-zero on ERRORS so the caller can fix and re-run
(produce -> validate -> fix -> re-validate).

Checks:
  ERROR   - Unfilled template tags ({{ ... }} / {% ... %}) anywhere, incl. tables.
  ERROR   - Leftover placeholder text (TBD, TODO, FIXME, XXX, lorem/ipsum, PLACEHOLDER).
  WARNING - No heading-styled paragraphs (structure/TOC will be empty; fake/bold
            headings do not count).
  WARNING - Inline images with no alt text (accessibility).
  WARNING - Page size is neither US Letter nor A4 (a common default-size mistake).

Usage:
    python scripts/validate_docx.py path/to/file.docx
"""
import json
import re
import sys
from pathlib import Path

from docx import Document

TAG = re.compile(r"{{.*?}}|{%.*?%}")
PLACEHOLDER = re.compile(
    r"\bTBD\b|\bTODO\b|\bFIXME\b|\bXXX+\b|\bPLACEHOLDER\b|lorem\s+ipsum",
    re.IGNORECASE,
)


def all_paragraphs_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            parts.extend(p.text for p in hf.paragraphs)
    return parts


def has_heading_style(doc):
    for p in doc.paragraphs:
        name = (p.style.name if p.style else "") or ""
        if name == "Title" or name.startswith("Heading"):
            return True
    return False


def images_missing_alt(doc):
    missing = 0
    for shape in doc.inline_shapes:
        try:
            descr = shape._inline.docPr.get("descr")
            title = shape._inline.docPr.get("title")
        except Exception:
            descr = title = None
        if not (descr or title):
            missing += 1
    return missing, len(doc.inline_shapes)


def page_size_ok(doc):
    for section in doc.sections:
        w = section.page_width.inches if section.page_width else 0
        h = section.page_height.inches if section.page_height else 0
        letter = abs(w - 8.5) < 0.1 and abs(h - 11) < 0.1
        a4 = abs(w - 8.27) < 0.15 and abs(h - 11.69) < 0.15
        if not (letter or a4):
            return False, round(w, 2), round(h, 2)
    return True, None, None


def main():
    if len(sys.argv) != 2:
        sys.exit("usage: python scripts/validate_docx.py <file.docx>")
    path = Path(sys.argv[1])
    if not path.exists():
        sys.exit(f"file not found: {path}")

    doc = Document(str(path))
    texts = all_paragraphs_text(doc)
    joined = "\n".join(texts)

    errors, warnings = [], []
    tags = sorted(set(TAG.findall(joined)))
    if tags:
        errors.append(f"Unfilled template tags remain: {tags}")
    placeholders = sorted(set(m.group(0) for m in PLACEHOLDER.finditer(joined)))
    if placeholders:
        errors.append(f"Leftover placeholder text: {placeholders}")

    if not has_heading_style(doc):
        warnings.append("No heading-styled paragraphs found; TOC/navigation will be "
                        "empty (bold text is not a heading).")
    missing_alt, n_images = images_missing_alt(doc)
    if missing_alt:
        warnings.append(f"{missing_alt}/{n_images} inline image(s) have no alt text.")
    ok, w, h = page_size_ok(doc)
    if not ok:
        warnings.append(f"Page size {w}x{h} in is neither US Letter (8.5x11) nor A4.")

    report = {
        "file": str(path),
        "status": "OK" if not errors else "FAILED",
        "errors": errors,
        "warnings": warnings,
        "has_headings": has_heading_style(doc),
    }
    print(json.dumps(report, indent=2, ensure_ascii=False))
    sys.exit(0 if not errors else 1)


if __name__ == "__main__":
    main()
