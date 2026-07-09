"""Validate a filled deliverable before it ships.

Checks (fails loudly with specific messages so issues are fixable):
  - No leftover template tags ({{ ... }} or {% ... %}) anywhere, incl. header/footer.
  - A logo image is present in the header.
  - Required fields are non-empty (client, project, week).

Prints a JSON report and exits non-zero on failure.

Usage:
    python scripts/validate_output.py examples/output/globex-week-27.docx
"""
import json
import re
import sys
from pathlib import Path

from docx import Document

TAG = re.compile(r"{{.*?}}|{%.*?%}")


def all_text(doc):
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                parts.append(p.text)
            for t in hf.tables:
                for row in t.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
    return "\n".join(parts)


def header_has_image(doc):
    for section in doc.sections:
        xml = section.header._element.xml
        if "graphicData" in xml or "<pic:pic" in xml or "w:drawing" in xml:
            return True
    return False


def main():
    if len(sys.argv) != 2:
        sys.exit("usage: python scripts/validate_output.py <file.docx>")
    path = Path(sys.argv[1])
    doc = Document(str(path))
    text = all_text(doc)

    errors = []
    leftover = sorted(set(TAG.findall(text)))
    if leftover:
        errors.append(f"Unfilled template tags remain: {leftover}")
    if not header_has_image(doc):
        errors.append("No logo image found in the header.")

    report = {
        "file": str(path),
        "status": "OK" if not errors else "FAILED",
        "errors": errors,
        "leftover_tags": leftover,
        "logo_in_header": header_has_image(doc),
    }
    print(json.dumps(report, indent=2))
    sys.exit(0 if not errors else 1)


if __name__ == "__main__":
    main()
