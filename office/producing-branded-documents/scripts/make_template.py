"""Generate the standard weekly-project-update Word template.

Run ONCE to (re)create assets/templates/weekly-project-update.docx from the brand
spec below. The output is a docxtpl template: the header, footer, logo frame,
colors, and fonts are the *invariant* (same for every team); the body contains
{{ jinja }} placeholders that fill_docx.py replaces per project/client.

Building the template programmatically (rather than typing tags in Word) keeps
each jinja tag in a single run, so docxtpl always parses it cleanly.

Usage:
    python scripts/make_template.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

# --- Brand spec (the invariant; edit here to rebrand everything) ----------------
BRAND = {
    "primary": RGBColor(0x1F, 0x3A, 0x5F),   # navy — headings, header/footer rule
    "accent": RGBColor(0x2E, 0x86, 0xAB),    # blue — table header fill text
    "text": RGBColor(0x22, 0x22, 0x22),
    "muted": RGBColor(0x66, 0x66, 0x66),
    "body_font": "Calibri",
    "heading_font": "Calibri",
    "logo_height_in": 0.5,                   # fixed logo height in the header
}
BASE = Path(__file__).resolve().parents[1]
OUT = BASE / "assets" / "templates" / "weekly-project-update.docx"


def set_font(run, *, name=None, size=None, bold=None, color=None):
    if name:
        run.font.name = name
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    """Insert a live PAGE number field (python-docx has no direct API)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def build():
    doc = Document()

    # Base styles
    normal = doc.styles["Normal"]
    normal.font.name = BRAND["body_font"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BRAND["text"]

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # --- Header: logo frame (left) + confidential label (right) ----------------
    header = section.header
    htab = header.add_table(rows=1, cols=2, width=Inches(6.7))
    htab.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = htab.rows[0].cells
    # Logo placeholder: docxtpl swaps {{ logo }} for the client's InlineImage.
    lp = left.paragraphs[0]
    lp.add_run("{{ logo }}")
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = rp.add_run("CONFIDENTIAL")
    set_font(r, size=8, bold=True, color=BRAND["muted"])
    rp.add_run("\n")
    r2 = rp.add_run("{{ client_name }}")
    set_font(r2, size=8, color=BRAND["muted"])

    # --- Footer: page number + confidentiality ---------------------------------
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Confidential — {{ company_name }}   |   Page ")
    set_font(r, size=8, color=BRAND["muted"])
    add_page_field(fp)

    # --- Title -----------------------------------------------------------------
    title = doc.add_paragraph()
    tr = title.add_run("Weekly Project Update")
    set_font(tr, name=BRAND["heading_font"], size=22, bold=True, color=BRAND["primary"])
    sub = doc.add_paragraph()
    sr = sub.add_run("{{ project_name }}  ·  Week {{ week }}  ·  {{ report_date }}")
    set_font(sr, size=11, color=BRAND["muted"])

    # --- Metadata table --------------------------------------------------------
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    labels = ["Client", "Project", "Prepared by", "Overall status"]
    values = ["{{ client_name }}", "{{ project_name }}",
              "{{ prepared_by }}", "{{ overall_status }}"]
    for i, lab in enumerate(labels):
        c = meta.rows[0].cells[i]
        shade_cell(c, "1F3A5F")
        run = c.paragraphs[0].add_run(lab)
        set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, val in enumerate(values):
        c = meta.rows[1].cells[i]
        c.paragraphs[0].add_run(val)
    doc.add_paragraph()

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, name=BRAND["heading_font"], size=13, bold=True, color=BRAND["primary"])
        # bottom rule
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "2E86AB")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def bulleted_loop(list_name):
        """A docxtpl paragraph-loop rendering each item as a bullet."""
        p = doc.add_paragraph()
        p.add_run("{%%p for item in %s %%}" % list_name)
        b = doc.add_paragraph(style="List Bullet")
        b.add_run("{{ item }}")
        e = doc.add_paragraph()
        e.add_run("{%p endfor %}")

    heading("Summary")
    doc.add_paragraph().add_run("{{ summary }}")

    heading("Accomplishments this week")
    bulleted_loop("accomplishments")

    heading("In progress")
    bulleted_loop("in_progress")

    heading("Blockers & risks")
    bulleted_loop("blockers")

    heading("Upcoming next week")
    bulleted_loop("upcoming")

    heading("Notes")
    doc.add_paragraph().add_run("{{ notes }}")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Template written: {OUT}")


if __name__ == "__main__":
    build()
