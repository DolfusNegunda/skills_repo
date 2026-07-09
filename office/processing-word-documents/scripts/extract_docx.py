"""Extract a .docx into clean Markdown + structured tables, with a fidelity report.

Deterministic ingestion so a model works on faithful, structured content instead of a
flattened text dump. Preserves the heading hierarchy (Word heading styles -> #/##/###),
renders tables as Markdown, walks paragraphs and tables in document order, and reports
the easily-missed parts (tracked changes, comments, footnotes, images, headers/footers,
merged table cells) in a `fidelity` block so nothing silently disappears.

This is an EXTRACTOR (ingest + fidelity self-check), not a validate->fix loop.
Legacy `.doc` (binary) is not supported here — convert to `.docx` first.

Usage:
    python scripts/extract_docx.py path/to/file.docx
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def heading_level(style_name):
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name == "title":
        return 1
    if name.startswith("heading "):
        try:
            return min(int(name.split()[1]), 6)
        except (ValueError, IndexError):
            return None
    return None


def escape_pipes(text):
    return text.replace("|", "\\|").replace("\n", " ").strip()


def table_has_merges(table):
    xml = table._tbl.xml
    return ("gridSpan" in xml) or ("vMerge" in xml)


def render_table_md(table):
    rows = []
    for row in table.rows:
        rows.append([escape_pipes(c.text) for c in row.cells])
    if not rows:
        return "", []
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join(["---"] * width) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines), rows


def body_blocks(doc):
    """Yield ('p', Paragraph) / ('tbl', Table) in true document order."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "tbl", Table(child, doc)


def hf_has_content(hf):
    """True only if a header/footer has real content (text, table, or image).

    python-docx always returns >=1 (usually empty) paragraph, so presence of a
    paragraph object is not evidence of content.
    """
    if any(p.text.strip() for p in hf.paragraphs):
        return True
    if hf.tables:
        return True
    xml = hf._element.xml
    return ("w:drawing" in xml) or ("pic:pic" in xml)


def main():
    if len(sys.argv) != 2:
        sys.exit("usage: python scripts/extract_docx.py <file.docx>")
    path = Path(sys.argv[1])
    if not path.exists():
        sys.exit(f"file not found: {path}")
    if path.suffix.lower() == ".doc":
        sys.exit("legacy .doc is not supported; convert to .docx first (LibreOffice/pandoc).")

    doc = Document(str(path))

    md_parts, tables, warnings = [], [], []
    table_idx = 0
    for kind, block in body_blocks(doc):
        if kind == "p":
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style else ""
            lvl = heading_level(style)
            if lvl:
                md_parts.append(f"{'#' * lvl} {text}")
            elif style and style.lower().startswith("list"):
                md_parts.append(f"- {text}")
            else:
                md_parts.append(text)
        else:  # table
            md, rows = render_table_md(block)
            merged = table_has_merges(block)
            tables.append({"index": table_idx, "n_rows": len(rows),
                           "n_cols": len(rows[0]) if rows else 0,
                           "merged_cells": merged, "rows": rows})
            md_parts.append(md)
            if merged:
                warnings.append(
                    f"Table #{table_idx} has merged/spanned cells; text repeats across "
                    "the span and rows may misalign — verify against the source."
                )
            table_idx += 1

    # Fidelity: detect the easily-missed parts via *references in the body*, not
    # part existence. footnotes.xml/comments.xml carry default entries (separators,
    # ids -1/0) even with zero real footnotes/comments, so a part-existence check
    # false-positives on ordinary docs. A reference tag appears only when one is used.
    tags = {el.tag for el in doc.element.body.iter()}
    has_tracked = qn("w:ins") in tags or qn("w:del") in tags
    has_footnotes = qn("w:footnoteReference") in tags
    has_comments = qn("w:commentReference") in tags
    n_images = len(doc.inline_shapes)
    has_header = any(hf_has_content(s.header) for s in doc.sections)
    has_footer = any(hf_has_content(s.footer) for s in doc.sections)

    if has_tracked:
        warnings.append("Tracked changes present: output reflects current text; "
                        "decide accepted vs. original explicitly.")
    if has_comments:
        warnings.append("Comments present but not inlined; extract separately if needed.")
    if has_footnotes:
        warnings.append("Footnotes present but not inlined; extract separately if needed.")
    if n_images:
        warnings.append(f"{n_images} inline image(s) not extracted as text.")

    result = {
        "file": str(path),
        "format": "docx",
        "markdown": "\n\n".join(md_parts),
        "tables": tables,
        "fidelity": {
            "n_tables": len(tables),
            "n_inline_images": n_images,
            "has_tracked_changes": has_tracked,
            "has_comments": has_comments,
            "has_footnotes": has_footnotes,
            "has_headers": bool(has_header),
            "has_footers": bool(has_footer),
            "warnings": warnings,
        },
    }
    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0)


if __name__ == "__main__":
    main()
